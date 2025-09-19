# -*- coding: utf-8 -*-
# pages/02_Paciente_Detalhe.py

import streamlit as st
import pandas as pd
import numpy as np
from datetime import datetime, date
from io import BytesIO

# ---- utils do app
from utils_casulo import connect, read_ws, append_rows, new_id
try:
    from utils_casulo import default_profissional
except Exception:
    def default_profissional():
        return st.secrets.get("DEFAULT_PROFISSIONAL", "Fernanda")

# ---- telegram (opcional, mas já preparado)
try:
    from utils_telegram import tg_send_document
except Exception:
    def tg_send_document(*args, **kwargs):
        return False, "utils_telegram não encontrado."


st.set_page_config(page_title="Casulo — Paciente", page_icon="📄", layout="wide")
st.title("📄 Detalhe do Paciente")


# ---------------- helpers ----------------
def to_date(s):
    if s is None:
        return None
    s = str(s).strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            pass
    return None

def brl(v: float) -> str:
    try:
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X",".")
    except Exception:
        return "R$ 0,00"

def to_float(x) -> float:
    try:
        s = str(x).strip().replace("R$", "").replace(" ", "")
        if s.count(",")==1 and s.count(".")>=1:
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", ".")
        return float(s)
    except Exception:
        return 0.0

def _clean(df: pd.DataFrame, cols: list[str] | None = None) -> pd.DataFrame:
    """Higieniza NaN -> '', tira 'nan' string e trim."""
    if df is None or df.empty:
        return df
    df = df.replace({np.nan: ""})
    if cols:
        for c in cols:
            if c in df.columns:
                df[c] = df[c].astype(str).replace("nan", "").str.strip()
    return df


# ---------------- conexão + colunas ----------------
ss = connect()

PAC_COLS = ["PacienteID","Nome","DataNascimento","Responsavel","Telefone","Email",
            "Diagnostico","Convenio","Status","Prioridade","FotoURL","Observacoes"]

SES_COLS = ["SessaoID","PacienteID","Data","HoraInicio","HoraFim","Profissional","Status",
            "Tipo","ObjetivosTrabalhados","Observacoes","AnexosURL"]

# nova aba de relatórios
REL_COLS = ["RelatorioID","PacienteID","Data","Tipo","Titulo","Texto","Autor","Privado","AnexosURL"]

PAG_COLS = ["PagamentoID","PacienteID","Data","Forma","Bruto","Liquido","TaxaValor","TaxaPct","Referencia","Obs","ReciboURL"]

df_pac, _ = read_ws(ss, "Pacientes",  PAC_COLS)
df_ses, _ = read_ws(ss, "Sessoes",    SES_COLS)
df_rel, ws_rel = read_ws(ss, "Relatorios", REL_COLS)           # <— cria/garante aba Relatorios
df_pag, _ = read_ws(ss, "Pagamentos", PAG_COLS)

# ---- limpeza de dados (elimina nan/N a N e espaços) ----
df_pac = _clean(df_pac, ["Nome","FotoURL","Responsavel","Telefone","Diagnostico","Convenio","Status","Prioridade","Observacoes"])
df_ses = _clean(df_ses, ["Data","HoraInicio","HoraFim","Profissional","Status","Tipo","ObjetivosTrabalhados","Observacoes","AnexosURL"])
df_rel = _clean(df_rel, ["RelatorioID","PacienteID","Data","Tipo","Titulo","Texto","Autor","Privado","AnexosURL"])
df_pag = _clean(df_pag, ["Data","Forma","Bruto","Liquido","TaxaValor","TaxaPct","Referencia","Obs","ReciboURL"])

# normalizações numéricas e datas auxiliares
if not df_ses.empty:
    df_ses["__dt"] = df_ses["Data"].apply(to_date)
if not df_pag.empty:
    df_pag["__dt"] = df_pag["Data"].apply(to_date)
    df_pag["__brl"] = df_pag["Liquido"].apply(to_float)


# ---------------- seleção por NOME ----------------
nomes = [""] + sorted(df_pac["Nome"].astype(str).str.strip().unique().tolist())
nome_sel = st.selectbox("Paciente", nomes, index=0, placeholder="Digite o nome…")
if not nome_sel:
    st.info("Selecione um paciente pelo nome.")
    st.stop()

p_row = df_pac[df_pac["Nome"].astype(str).str.strip() == nome_sel].head(1)
if p_row.empty:
    st.warning("Paciente não encontrado.")
    st.stop()
p = p_row.iloc[0]
pid = str(p["PacienteID"])


# ---------------- Header do paciente ----------------
colA, colB = st.columns([1,3])
with colA:
    foto = str(p.get("FotoURL","") or "").strip()
    if foto:
        st.image(foto, caption=p.get("Nome",""), width=220)
with colB:
    st.markdown(f"## {p.get('Nome','')}")
    st.write(f"**Responsável:** {p.get('Responsavel','-')}  |  **Telefone:** {p.get('Telefone','-')}")
    st.write(f"**Diagnóstico:** {p.get('Diagnostico','-')}")
    st.write(f"**Convênio:** {p.get('Convenio','-')}  |  **Status:** {p.get('Status','-')}  |  **Prioridade:** {p.get('Prioridade','-')}")
    if str(p.get("Observacoes","")).strip():
        st.caption(p.get("Observacoes",""))


# ---------------- KPIs do paciente ----------------
ses_cli = df_ses[df_ses["PacienteID"].astype(str) == pid] if not df_ses.empty else df_ses.iloc[0:0]
rel_cli = df_rel[df_rel["PacienteID"].astype(str) == pid] if not df_rel.empty else df_rel.iloc[0:0]
pag_cli = df_pag[df_pag["PacienteID"].astype(str) == pid] if not df_pag.empty else df_pag.iloc[0:0]

k1, k2, k3, k4 = st.columns(4)
k1.metric("Sessões", int(len(ses_cli)))
k2.metric("Recebido (líquido)", brl(float(pag_cli["__brl"].sum()) if not pag_cli.empty else 0.0))
k3.metric("Relatórios", int(len(rel_cli)))
ultima_data = None
if not ses_cli.empty and "__dt" in ses_cli:
    _ok = ses_cli.dropna(subset=["__dt"])
    if not _ok.empty:
        ultima_data = _ok["__dt"].max().strftime("%d/%m/%Y")
k4.metric("Última sessão", ultima_data or "—")

st.divider()


# ---------------- Tabs ----------------
tab_geral, tab_rel, tab_ses, tab_fin, tab_docs = st.tabs(
    ["👀 Visão geral", "🧾 Relatórios", "📝 Sessões", "💰 Financeiro", "📎 Documentos"]
)


# =============== Visão Geral: timeline simples ===============
with tab_geral:
    st.subheader("Linha do tempo")
    eventos = []

    # sessões
    if not ses_cli.empty:
        for _, r in ses_cli.iterrows():
            d = to_date(r.get("Data",""))
            if not d:
                continue
            eventos.append({
                "dt": d,
                "tipo": "Sessão",
                "titulo": str(r.get("Tipo","Terapia")),
                "sub": f"{r.get('HoraInicio','--')}-{r.get('HoraFim','')} · {r.get('Profissional','') or default_profissional()} · {r.get('Status','Agendada')}"
            })

    # relatórios
    if not rel_cli.empty:
        for _, r in rel_cli.iterrows():
            d = to_date(r.get("Data",""))
            if not d:
                continue
            eventos.append({
                "dt": d,
                "tipo": "Relatório",
                "titulo": str(r.get("Titulo","(sem título)")),
                "sub": f"{r.get('Tipo','-')} · Autor: {r.get('Autor','') or default_profissional()}"
            })

    # ordenar do mais antigo -> mais novo
    eventos = sorted(eventos, key=lambda x: x["dt"])

    if not eventos:
        st.info("Sem eventos para este paciente.")
    else:
        for e in eventos:
            st.markdown(
                f"**{e['dt'].strftime('%d/%m/%Y')}** — *{e['tipo']}*\n\n"
                f"**{e['titulo']}**  \n"
                f"<span style='color:#94a3b8'>{e['sub']}</span>",
                unsafe_allow_html=True
            )
            st.markdown("---")


# =============== RELATÓRIOS ===============
with tab_rel:
    st.subheader("Relatórios do paciente")

    # filtros leves
    colR1, colR2, colR3 = st.columns([1,1,1])
    with colR1:
        tipo_f = st.selectbox("Tipo", ["(todos)","Anamnese","Evolução","Avaliação","Alta","Outro"], index=0)
    with colR2:
        di = st.date_input("De", value=None)
    with colR3:
        df_ = st.date_input("Até", value=None)

    lst = rel_cli.copy()
    if not lst.empty:
        lst["__dt"] = lst["Data"].apply(to_date)
        if tipo_f != "(todos)":
            lst = lst[lst["Tipo"].astype(str) == tipo_f]
        if di:
            lst = lst[lst["__dt"] >= di]
        if df_:
            lst = lst[lst["__dt"] <= df_]
        lst = lst.sort_values("__dt", ascending=True)

    ids_sel = st.multiselect(
        "Selecionar relatórios para exportar",
        options=(lst["RelatorioID"].astype(str).tolist() if not lst.empty else []),
        format_func=lambda rid: (
            f"{rid} · {lst.loc[lst['RelatorioID'].astype(str)==rid, 'Titulo'].values[0]}"
            if (not lst.empty and (lst['RelatorioID'].astype(str)==rid).any()) else rid
        )
    )

    # listagem com possibilidade de excluir
    if lst.empty:
        st.info("Sem relatórios no filtro atual.")
    else:
        for _, r in lst.iterrows():
            rid = str(r.get("RelatorioID"))
            with st.expander(f"🧾 {r.get('Data','--')} — {r.get('Tipo','-')} · {r.get('Titulo','(sem título)')}", expanded=False):
                st.markdown(f"**Autor:** {r.get('Autor','-')}  ·  **Privado:** {r.get('Privado','Não')}")
                st.write(r.get("Texto",""))
                anex = str(r.get("AnexosURL","") or "").strip()
                if anex:
                    st.markdown(f"🔗 **Anexos:** {anex}")
                cA, cB = st.columns([1,3])
                with cA:
                    conf = st.checkbox(f"Confirmar exclusão ({rid})", key=f"conf_rel_{rid}")
                with cB:
                    if st.button("🗑️ Excluir", key=f"del_rel_{rid}") and conf:
                        try:
                            idx = df_rel.index[df_rel["RelatorioID"].astype(str)==rid]
                            if len(idx)>0:
                                row = int(idx[0]) + 2
                                ws_rel.delete_rows(row)
                                st.success(f"Relatório {rid} excluído.")
                                st.cache_data.clear()
                                st.rerun()
                        except Exception as e:
                            st.error(f"Erro ao excluir: {e}")

    # -------- exportação / envio --------
    def _coletar_md(ids: list[str]) -> str:
        base = rel_cli if not rel_cli.empty else df_rel[df_rel["PacienteID"].astype(str)==pid]
        # limpeza adicional de strings
        if not base.empty:
            for c in ["Data","Tipo","Titulo","Texto","Autor","Privado","AnexosURL"]:
                if c in base.columns:
                    base[c] = base[c].astype(str).replace("nan","").str.strip()

        blocos = [f"# Relatórios — {nome_sel}", ""]
        for rid in ids:
            r = base[base["RelatorioID"].astype(str)==rid]
            if r.empty:
                continue
            rr = r.iloc[0]
            data_fmt = to_date(rr.get("Data",""))
            data_fmt = data_fmt.strftime("%d/%m/%Y") if data_fmt else "--"
            titulo = rr.get("Titulo","") or "(sem título)"
            tipo   = rr.get("Tipo","-") or "-"
            autor  = rr.get("Autor","") or default_profissional()
            priv   = rr.get("Privado","Não") or "Não"
            texto  = rr.get("Texto","") or ""
            anexos = rr.get("AnexosURL","") or ""

            blocos += [
                f"## {data_fmt} — {tipo} · {titulo}",
                f"**Autor:** {autor}  ·  **Privado:** {priv}",
                "",
                texto,
                "",
            ]
            if anexos:
                blocos += [f"**Anexos:** {anexos}", ""]
        return "\n".join(blocos)

    def _gerar_pdf(md_text: str) -> bytes:
        # PDF PRO com cabeçalho, rodapé e tipografia melhor (ReportLab)
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer)
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import mm

        # fontes (usa padrão caso não tenha TTF)
        try:
            pdfmetrics.registerFont(TTFont("Inter", "Inter-Regular.ttf"))
            pdfmetrics.registerFont(TTFont("Inter-Bold", "Inter-Bold.ttf"))
            base_font = "Inter"
        except Exception:
            base_font = "Helvetica"

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=18*mm, rightMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle('h1', parent=styles['Heading1'], alignment=TA_LEFT, spaceAfter=8, fontName=base_font)
        h2 = ParagraphStyle('h2', parent=styles['Heading2'], alignment=TA_LEFT, spaceAfter=6, fontName=base_font)
        p  = ParagraphStyle('p',  parent=styles['BodyText'],  leading=15, fontName=base_font)

        def _header_footer(canvas, doc_):
            canvas.saveState()
            canvas.setFont(base_font, 10)
            canvas.drawString(18*mm, A4[1]-12*mm, f"Relatórios — {nome_sel}")
            canvas.drawRightString(A4[0]-18*mm, 12*mm, f"Página {doc_.page}")
            canvas.restoreState()

        story = []
        # capa simples
        story.append(Paragraph(f"Relatórios — {nome_sel}", h1))
        story.append(Spacer(1, 8))

        # parse md simples
        for ln in md_text.splitlines():
            if ln.startswith("# "):
                continue  # capa já feita
            elif ln.startswith("## "):
                story.append(Spacer(1, 6))
                story.append(Paragraph(ln[3:], h2))
            elif ln.strip() == "":
                story.append(Spacer(1, 4))
            else:
                story.append(Paragraph(ln, p))

        doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
        return buf.getvalue()

    def _gerar_docx(md_text: str) -> bytes:
        from docx import Document
        doc = Document()
        for ln in md_text.splitlines():
            if ln.startswith("# "):   doc.add_heading(ln[2:], level=1)
            elif ln.startswith("## "): doc.add_heading(ln[3:], level=2)
            else:                      doc.add_paragraph(ln)
        buf = BytesIO(); doc.save(buf); return buf.getvalue()

    # leitura rápida (HTML) — opcional
    _preview_html = st.toggle("Leitura rápida no app (HTML)", value=False)
    if _preview_html and ids_sel:
        md_str = _coletar_md(ids_sel)
        try:
            import markdown
            html = markdown.markdown(md_str)
            st.markdown(html, unsafe_allow_html=True)
        except Exception:
            st.code(md_str, language="markdown")

    colx1, colx2, colx3, colx4, colx5 = st.columns([1,1,1,2,2])
    with colx1:
        if st.button("⬇️ MD") and ids_sel:
            md = _coletar_md(ids_sel).encode("utf-8")
            st.download_button("Baixar .md", data=md, file_name=f"relatorios_{pid}.md",
                               mime="text/markdown", use_container_width=True)
    with colx2:
        if st.button("⬇️ PDF") and ids_sel:
            pdf_bytes = _gerar_pdf(_coletar_md(ids_sel))
            st.download_button("Baixar PDF", data=pdf_bytes, file_name=f"relatorios_{pid}.pdf",
                               mime="application/pdf", use_container_width=True)
    with colx3:
        if st.button("⬇️ DOCX") and ids_sel:
            docx_bytes = _gerar_docx(_coletar_md(ids_sel))
            st.download_button("Baixar DOCX", data=docx_bytes,
                               file_name=f"relatorios_{pid}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True)
    with colx4:
        if st.button("📤 Enviar PDF ao Telegram") and ids_sel:
            pdf_bytes = _gerar_pdf(_coletar_md(ids_sel))
            ok, err = tg_send_document(
                data=pdf_bytes,
                filename=f"relatorios_{pid}.pdf",
                mime="application/pdf",
                caption=f"Relatórios — {nome_sel}"
            )
            if ok: st.success("Enviado ao Telegram ✅")
            else:  st.error(f"Falhou ao enviar: {err}")
    with colx5:
        if st.button("👁️ Pré-visualizar no app") and ids_sel:
            import base64
            pdf_bytes = _gerar_pdf(_coletar_md(ids_sel))
            b64 = base64.b64encode(pdf_bytes).decode()
            st.components.v1.html(
                f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="720px" style="border:none;"></iframe>',
                height=740
            )

    st.markdown("---")
    st.subheader("Novo relatório")

    with st.form("novo_relatorio"):
        c1, c2 = st.columns([2,1])
        with c1:
            titulo = st.text_input("Título", "")
            texto  = st.text_area("Texto (anotações, evolução, anamnese…)", height=220)
        with c2:
            data_r = st.date_input("Data", value=date.today())
            tipo   = st.selectbox("Tipo", ["Evolução","Anamnese","Avaliação","Alta","Outro"], index=0)
            autor  = st.text_input("Autor", default_profissional())
            privado = st.selectbox("Privado", ["Não","Sim"], index=0)
            anexos = st.text_input("AnexosURL (opcional)", "")

        ok = st.form_submit_button("Salvar relatório")
        if ok:
            rid = new_id("R")
            append_rows(ws_rel, [{
                "RelatorioID": rid,
                "PacienteID":  pid,
                "Data":        data_r.strftime("%d/%m/%Y"),
                "Tipo":        tipo,
                "Titulo":      (titulo or "").strip() or "(sem título)",
                "Texto":       (texto or "").strip(),
                "Autor":       (autor or "").strip() or default_profissional(),
                "Privado":     privado,
                "AnexosURL":   (anexos or "").strip()
            }], default_headers=REL_COLS)
            st.success(f"Relatório salvo ({rid}).")
            st.cache_data.clear()
            st.rerun()


# =============== Sessões ===============
with tab_ses:
    st.subheader("Sessões do paciente")
    if ses_cli.empty:
        st.info("Sem sessões registradas.")
    else:
        cols = ["Data","HoraInicio","HoraFim","Profissional","Status","Tipo","ObjetivosTrabalhados","Observacoes"]
        cols = [c for c in cols if c in ses_cli.columns]
        # ordena por data/hora crescente
        tmp = ses_cli.copy()
        tmp["__dt"] = tmp["Data"].apply(to_date)
        tmp = tmp.sort_values(["__dt","HoraInicio"], ascending=True)
        st.dataframe(tmp[cols], use_container_width=True, hide_index=True)


# =============== Financeiro ===============
with tab_fin:
    st.subheader("Recebimentos do paciente")
    if pag_cli.empty:
        st.info("Sem pagamentos registrados.")
    else:
        viz = pag_cli.copy()
        viz["Liquido"] = viz["__brl"].apply(brl)
        cols = ["Data","Forma","Bruto","Liquido","TaxaValor","Referencia","Obs"]
        cols = [c for c in cols if c in viz.columns]
        st.dataframe(viz[cols], use_container_width=True, hide_index=True)
        st.metric("Total líquido", brl(float(pag_cli["__brl"].sum())))


# =============== Documentos & Anexos ===============
with tab_docs:
    st.subheader("Documentos & anexos")

    st.info(
        "Para anexar um link/arquivo, use o campo **AnexosURL** ao criar um Relatório na aba **Relatórios**.\n"
        "Se quiser uma aba exclusiva na planilha (ex.: 'Documentos' com colunas: PacienteID, Título, Data, URL, Obs), "
        "dá para integrar depois."
    )

    docs = rel_cli.copy() if not rel_cli.empty else df_rel[df_rel["PacienteID"].astype(str) == pid]
    if docs.empty:
        st.caption("Sem anexos para este paciente.")
    else:
        docs["__url"] = docs["AnexosURL"].astype(str).str.strip()
        docs = docs[docs["__url"] != ""]
        if docs.empty:
            st.caption("Sem anexos com URL preenchido.")
        else:
            def _nice_title(row):
                t = (row.get("Titulo") or "").strip()
                if not t:
                    t = (row.get("Tipo") or "Documento").strip()
                return t or "Documento"

            docs["__titulo"] = docs.apply(_nice_title, axis=1)
            docs["__data"] = docs["Data"].apply(lambda s: to_date(s).strftime("%d/%m/%Y") if to_date(s) else "--")

            for _, r in docs.sort_values("__data", ascending=True).iterrows():
                autor = (r.get("Autor","") or "").strip() or nome_sel
                st.markdown(f"- **{r['__data']}** — {autor} · [{r['__titulo']}]({r['__url']})")
