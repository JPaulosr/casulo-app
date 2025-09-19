# -*- coding: utf-8 -*-
# pages/02_Paciente_Detalhe.py

import io
import os
import base64
import requests
from datetime import datetime, date
import pandas as pd
import numpy as np
import streamlit as st

from utils_casulo import connect, read_ws, append_rows, new_id  # usa o appender SEGURO

# =========================
# Config & constantes
# =========================
st.set_page_config(page_title="Casulo — Paciente", page_icon="📄", layout="wide")
st.title("📄 Detalhe do Paciente")

CLINIC_NAME = "Espaço Terapêutico Casulo"

# =========================
# Telegram (compatível com suas chaves)
# =========================
# Fallback também lê variáveis de ambiente para uso local
TELEGRAM_TOKEN_FALLBACK = os.getenv("TELEGRAM_TOKEN", "") or os.getenv("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHATID_FALLBACK = (
    os.getenv("TELEGRAM_CHAT_ID", "")
    or os.getenv("TELEGRAM_CHAT_ID_CASULO", "")
    or os.getenv("TELEGRAM_CHAT_ID_PADRAO", "")
)

def _tg_token() -> str:
    try:
        for k in ("TELEGRAM_TOKEN", "TELEGRAM_BOT_TOKEN"):
            v = (st.secrets.get(k, "") or "").strip()
            if v:
                return v
    except Exception:
        pass
    return (TELEGRAM_TOKEN_FALLBACK or "").strip()

def _tg_chat_id() -> str:
    try:
        for k in ("TELEGRAM_CHAT_ID", "TELEGRAM_CHAT_ID_CASULO", "TELEGRAM_CHAT_ID_PADRAO"):
            v = (st.secrets.get(k, "") or "").strip()
            if v:
                return v
    except Exception:
        pass
    return (TELEGRAM_CHATID_FALLBACK or "").strip()

def tg_send_pdf(file_bytes: bytes, filename: str, caption: str = "") -> tuple[bool, str]:
    token = _tg_token()
    chat_id = _tg_chat_id()
    if not token or not chat_id:
        return False, f"Secrets ausentes. Token? {'OK' if bool(token) else 'NÃO'} | ChatID? {'OK' if bool(chat_id) else 'NÃO'}"
    try:
        url = f"https://api.telegram.org/bot{token}/sendDocument"
        files = {"document": (filename, file_bytes, "application/pdf")}
        data = {"chat_id": chat_id, "caption": (caption or "")[:1024]}  # segurança
        r = requests.post(url, data=data, files=files, timeout=60)
        ok = r.ok and r.json().get("ok")
        return (bool(ok), "" if ok else f"HTTP {r.status_code}: {r.text}")
    except Exception as e:
        return False, f"Erro de rede: {e}"

# =========================
# Helpers
# =========================
DATA_FMT = "%d/%m/%Y"

def to_date(s):
    if s is None: return None
    s = str(s).strip()
    for fmt in ("%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"):
        try:
            return datetime.strptime(s, fmt).date()
        except Exception:
            pass
    return None

def brl(v: float) -> str:
    try:
        return f"R$ {float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X",".")
    except Exception:
        return "R$ 0,00"

def _clean(df: pd.DataFrame, cols: list[str] | None = None) -> pd.DataFrame:
    """Higieniza NaN -> '', remove 'nan' textual, trim."""
    if df is None or df.empty:
        return df
    df = df.replace({np.nan: ""})
    if cols:
        for c in cols:
            if c in df.columns:
                df[c] = df[c].astype(str).replace("nan", "").str.strip()
    return df

# =========================
# Leitura das planilhas
# =========================
ss = connect()

PAC_COLS = ["PacienteID","Nome","DataNascimento","Responsavel","Telefone","Email",
            "Diagnostico","Convenio","Status","Prioridade","FotoURL","Observacoes"]

SES_COLS = ["SessaoID","PacienteID","Data","HoraInicio","HoraFim",
            "Profissional","Status","Tipo","ObjetivosTrabalhados","Observacoes","AnexosURL"]

PAG_COLS = ["PagamentoID","PacienteID","Data","Forma","Bruto","Liquido",
            "TaxaValor","TaxaPct","Referencia","Obs","ReciboURL"]

# Relatórios do paciente (layout novo)
REL_COLS = ["RelatorioID","PacienteID","Data","Tipo","Titulo","Autor","Texto","ArquivoURL"]

df_pac, _ = read_ws(ss, "Pacientes",  PAC_COLS)
df_ses, _ = read_ws(ss, "Sessoes",    SES_COLS)
df_pag, _ = read_ws(ss, "Pagamentos", PAG_COLS)
df_rel, ws_rel = read_ws(ss, "Relatorios", REL_COLS)  # cria se não existe

# limpeza
df_pac = _clean(df_pac, ["Nome","FotoURL","Responsavel","Telefone","Diagnostico","Convenio","Status","Prioridade","Observacoes"])
df_ses = _clean(df_ses, ["Data","HoraInicio","HoraFim","Profissional","Status","Tipo","ObjetivosTrabalhados","Observacoes","AnexosURL"])
df_pag = _clean(df_pag, ["Data","Forma","Bruto","Liquido","TaxaValor","TaxaPct","Referencia","Obs","ReciboURL"])
df_rel = _clean(df_rel, ["RelatorioID","PacienteID","Data","Tipo","Titulo","Autor","Texto","ArquivoURL"])

# =========================
# Selecionar paciente por nome
# =========================
nomes = [""] + sorted(df_pac["Nome"].astype(str).str.strip().unique().tolist())
nome_sel = st.selectbox("Paciente", nomes, index=0, placeholder="Digite o nome…")

if not nome_sel:
    st.info("Selecione um paciente pelo nome.")
    st.stop()

p_row = df_pac[df_pac["Nome"].astype(str).str.strip() == nome_sel].head(1)
if p_row.empty:
    st.warning("Paciente não encontrado.")
    st.stop()

p = p_row.iloc[0]
pid = str(p["PacienteID"])

# =========================
# Header — foto + dados
# =========================
col1, col2 = st.columns([1,3])
with col1:
    foto = str(p.get("FotoURL","")).strip()
    if foto:
        st.image(foto, caption=p.get("Nome",""), width=260)
with col2:
    st.markdown(f"## {p.get('Nome','')}")
    st.write(f"**Responsável:** {p.get('Responsavel','-')}  |  **Telefone:** {p.get('Telefone','-')}")
    st.write(f"**Diagnóstico:** {p.get('Diagnostico','-')}")
    st.write(f"**Convênio:** {p.get('Convenio','-')}  |  **Status:** {p.get('Status','-')}  |  **Prioridade:** {p.get('Prioridade','-')}")
    st.caption(f"ID interno: {pid}")

st.divider()

# =========================
# KPIs do paciente
# =========================
df_ses_p = df_ses[df_ses["PacienteID"].astype(str) == pid].copy()
df_pag_p = df_pag[df_pag["PacienteID"].astype(str) == pid].copy()
df_rel_p = df_rel[df_rel["PacienteID"].astype(str) == pid].copy()

df_ses_p["__dt"] = df_ses_p.get("Data","").apply(to_date)
df_pag_p["__dt"] = df_pag_p.get("Data","").apply(to_date)
df_pag_p["__liq"] = pd.to_numeric(df_pag_p.get("Liquido",0), errors="coerce").fillna(0)

total_sessoes = int(len(df_ses_p))
realizadas = int((df_ses_p.get("Status","").astype(str).str.lower() == "realizada").sum())
recebido_liq = float(df_pag_p["__liq"].sum())
qtd_relatorios = int(len(df_rel_p))
ultima_sessao = df_ses_p["__dt"].dropna().max() if not df_ses_p.empty else None

k1, k2, k3, k4, k5 = st.columns(5)
k1.metric("Sessões", total_sessoes)
k2.metric("Realizadas", realizadas)
k3.metric("Recebido (líquido)", brl(recebido_liq))
k4.metric("Relatórios", qtd_relatorios)
k5.metric("Última sessão", ultima_sessao.strftime(DATA_FMT) if ultima_sessao else "-")

# =========================
# Abas
# =========================
tab_visao, tab_rel, tab_ses, tab_fin, tab_docs = st.tabs(
    ["👁️ Visão geral","📄 Relatórios","📝 Sessões","💰 Financeiro","📎 Documentos"]
)

# ---------- Visão geral ----------
with tab_visao:
    st.subheader("Linha do tempo")
    df_ses_p["__ord_h"] = pd.to_datetime(df_ses_p.get("HoraInicio",""), format="%H:%M", errors="coerce")
    df_ses_p = df_ses_p.sort_values(["__dt","__ord_h"], ascending=[True, True])

    s_status = df_ses_p.get("Status","").astype(str).str.lower()
    blocos = {
        "Agendadas/Confirmadas": df_ses_p[s_status.isin(["agendada","confirmada"])],
        "Realizadas": df_ses_p[s_status == "realizada"],
        "Faltas": df_ses_p[s_status == "falta"],
        "Canceladas": df_ses_p[s_status == "cancelada"],
    }

    for titulo, bloco in blocos.items():
        with st.expander(f"{titulo} ({len(bloco)})", expanded=(titulo=="Agendadas/Confirmadas")):
            if bloco.empty:
                st.info("Nada aqui.")
            else:
                for d, grupo in bloco.groupby("__dt"):
                    if pd.isna(d): continue
                    st.markdown(f"**{d.strftime(DATA_FMT)} — Sessão**")
                    for _, r in grupo.iterrows():
                        hi = str(r.get("HoraInicio","") or "").strip()
                        hf = str(r.get("HoraFim","") or "").strip()
                        prof = str(r.get("Profissional","") or "")
                        tipo = str(r.get("Tipo","Terapia") or "Terapia")
                        st.markdown(f"**{tipo}**  \n{hi}{('–'+hf) if hf else ''} · {prof}")

# ---------- util: montar MD e gerar PDF PRO ----------
def _compose_md(rows: pd.DataFrame, nome_paciente: str) -> str:
    parts = [f"# Relatórios — {nome_paciente}", ""]
    for _, r in rows.iterrows():
        d = to_date(r.get("Data"))
        dtxt = d.strftime(DATA_FMT) if d else "-"
        parts += [
            f"## {dtxt} — {str(r.get('Tipo','-'))} · {str(r.get('Titulo','(sem título)'))}",
            f"**Autor:** {str(r.get('Autor','-'))}",
            "",
            str(r.get("Texto","")).strip(),
            "",
        ]
        url = str(r.get("ArquivoURL","")).strip()
        if url:
            parts += [f"**Anexo:** {url}", ""]
    return "\n".join(parts)

def _gerar_pdf_pro(md_text: str, nome_paciente: str, clinic_name: str) -> bytes:
    """
    Gera PDF profissional com cabeçalho e rodapé (nome da clínica no rodapé).
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.units import mm

    # fonte (usa Inter se existir, senão Helvetica)
    try:
        pdfmetrics.registerFont(TTFont("Inter", "Inter-Regular.ttf"))
        pdfmetrics.registerFont(TTFont("Inter-Bold", "Inter-Bold.ttf"))
        base_font = "Inter"
    except Exception:
        base_font = "Helvetica"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18*mm, rightMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle('h1', parent=styles['Heading1'], alignment=TA_LEFT, spaceAfter=8, fontName=base_font)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'], alignment=TA_LEFT, spaceAfter=6, fontName=base_font)
    p  = ParagraphStyle('p',  parent=styles['BodyText'],  leading=15, fontName=base_font)

    def _header_footer(canvas, doc_):
        # Cabeçalho
        canvas.saveState()
        canvas.setFont(base_font, 10)
        canvas.drawString(18*mm, A4[1]-12*mm, f"Relatórios — {nome_paciente}")
        # Rodapé com nome da clínica e nº da página
        canvas.setFont(base_font, 10)
        canvas.drawString(18*mm, 12*mm, clinic_name)
        canvas.drawRightString(A4[0]-18*mm, 12*mm, f"Página {doc_.page}")
        canvas.restoreState()

    story = []
    # capa simples
    story.append(Paragraph(f"Relatórios — {nome_paciente}", h1))
    story.append(Spacer(1, 8))

    # parse simples do markdown
    for ln in md_text.splitlines():
        if ln.startswith("# "):
            continue  # capa já criada
        elif ln.startswith("## "):
            story.append(Spacer(1, 6))
            story.append(Paragraph(ln[3:], h2))
        elif ln.strip() == "":
            story.append(Spacer(1, 4))
        else:
            story.append(Paragraph(ln, p))

    doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
    return buf.getvalue()

def _preview_pdf_inline(pdf_bytes: bytes, filename: str):
    """Mostra PDF em iframe base64 + link 'abrir em nova aba'."""
    b64 = base64.b64encode(pdf_bytes).decode()
    dataurl = f"data:application/pdf;base64,{b64}"
    st.components.v1.html(
        f'<iframe src="{dataurl}" width="100%" height="720px" style="border:none;"></iframe>',
        height=740
    )
    st.markdown(f"[Abrir em nova aba]({dataurl})")

# ---------- Relatórios ----------
with tab_rel:
    st.subheader("Relatórios do paciente")

    # Filtros
    colf1, colf2, colf3 = st.columns([1,1,2])
    tipos = ["(todos)"] + sorted(df_rel_p.get("Tipo","").astype(str).str.strip().unique().tolist())
    with colf1:
        tipo_f = st.selectbox("Tipo", tipos, index=0)
    with colf2:
        de = st.date_input("De", value=None)
    with colf3:
        ate = st.date_input("Até", value=None)

    rel_vis = df_rel_p.copy()
    rel_vis["__dt"] = rel_vis.get("Data","").apply(to_date)
    if tipo_f != "(todos)":
        rel_vis = rel_vis[rel_vis.get("Tipo","").astype(str) == tipo_f]
    if de:
        rel_vis = rel_vis[rel_vis["__dt"] >= de]
    if ate:
        rel_vis = rel_vis[rel_vis["__dt"] <= ate]
    rel_vis = rel_vis.sort_values("__dt", ascending=True)

    # Lista compacta + seleção
    opts, labels = [], {}
    for _, r in rel_vis.iterrows():
        rid = str(r.get("RelatorioID",""))
        data_txt = r["__dt"].strftime(DATA_FMT) if pd.notna(r["__dt"]) else "-"
        titulo = str(r.get("Titulo","")).strip() or "(sem título)"
        lbl = f"🗂️ {data_txt} — {str(r.get('Tipo','')).strip() or '-'} · {titulo}"
        opts.append(rid)
        labels[rid] = lbl

    sel = st.multiselect("Selecionar relatórios para exportar", opts, format_func=lambda x: labels.get(x, x))

    colbtn1, colbtn2, colbtn3, colbtn4, colbtn5 = st.columns([1,1,1,2,2])
    with colbtn1:
        md_ok = st.button("⬇️ MD", use_container_width=True)
    with colbtn2:
        pdf_ok = st.button("⬇️ PDF", use_container_width=True)
    with colbtn3:
        docx_ok = st.button("⬇️ DOCX", use_container_width=True)
    with colbtn4:
        tg_ok = st.button("📤 Enviar PDF ao Telegram", use_container_width=True)
    with colbtn5:
        prev_ok = st.button("👁️ Pré-visualizar no app", use_container_width=True)

    rows_sel = rel_vis[rel_vis["RelatorioID"].astype(str).isin(sel)].copy()

    if md_ok:
        if rows_sel.empty:
            st.warning("Selecione ao menos um relatório.")
        else:
            md_txt = _compose_md(rows_sel, nome_sel)
            st.download_button("Baixar .md", data=md_txt.encode("utf-8"), file_name=f"relatorios_{pid}.md")

    if pdf_ok or tg_ok or prev_ok:
        if rows_sel.empty:
            st.warning("Selecione ao menos um relatório.")
        else:
            # PDF PRO com cabeçalho/rodapé (clínica no rodapé)
            md_txt = _compose_md(rows_sel, nome_sel)
            pdf_bytes = _gerar_pdf_pro(md_txt, nome_sel, CLINIC_NAME)

            if pdf_ok:
                st.download_button("Baixar PDF", data=pdf_bytes, file_name=f"relatorios_{pid}.pdf")
            if tg_ok:
                ok, err = tg_send_pdf(pdf_bytes, f"relatorios_{pid}.pdf", caption=f"Relatórios — {nome_sel}")
                if ok: st.success("Enviado ao Telegram ✅")
                else:  st.error(f"Falhou ao enviar: {err}")
            if prev_ok:
                _preview_pdf_inline(pdf_bytes, f"relatorios_{pid}.pdf")

    if docx_ok:
        if rows_sel.empty:
            st.warning("Selecione ao menos um relatório.")
        else:
            try:
                from docx import Document
                doc = Document()
                for _, r in rows_sel.iterrows():
                    d = to_date(r.get("Data"))
                    dtxt = d.strftime(DATA_FMT) if d else "-"
                    doc.add_heading(str(r.get("Titulo","(sem título)")), level=1)
                    doc.add_paragraph(f"{dtxt} • {str(r.get('Tipo','-'))} • {str(r.get('Autor','-'))}")
                    doc.add_paragraph(str(r.get("Texto","")))
                    url = str(r.get("ArquivoURL","")).strip()
                    if url:
                        doc.add_paragraph(f"Anexo: {url}")
                    doc.add_page_break()
                buf = io.BytesIO(); doc.save(buf)
                st.download_button("Baixar DOCX", data=buf.getvalue(), file_name=f"relatorios_{pid}.docx")
            except ImportError:
                st.error("Faltou a dependência `python-docx` para gerar DOCX.")

    # ---------- Diagnóstico Telegram (opcional) ----------
    with st.expander("🔧 Diagnóstico Telegram"):
        if st.button("Testar envio de mensagem"):
            tkn_ok = bool(_tg_token()); chat_ok = bool(_tg_chat_id())
            if not (tkn_ok and chat_ok):
                st.error(f"Token OK? {tkn_ok} | ChatID OK? {chat_ok}")
            else:
                try:
                    url = f"https://api.telegram.org/bot{_tg_token()}/sendMessage"
                    payload = {"chat_id": _tg_chat_id(), "text": "Teste ✅ Página Paciente_Detalhe ativa.", "parse_mode": "HTML"}
                    r = requests.post(url, json=payload, timeout=30)
                    if r.ok and r.json().get("ok"):
                        st.success("Mensagem enviada com sucesso ✅")
                    else:
                        st.error(f"Falhou: HTTP {r.status_code}: {r.text}")
                except Exception as e:
                    st.error(f"Erro: {e}")

    st.markdown("---")
    st.subheader("Novo relatório")

    # ====== Formulário de novo relatório ======
    with st.form("novo_rel"):
        c1, c2 = st.columns([2,1])
        with c1:
            titulo = st.text_input("Título", "")
        with c2:
            data_rel = st.date_input("Data", value=date.today(), format="YYYY/MM/DD")

        tipo = st.selectbox("Tipo", ["Evolução","Avaliação","Anamnese","Alta","Outro"], index=0)
        autor = st.text_input("Autor", "Fernanda")
        texto = st.text_area("Texto (anotações, evolução, anamnese...)", height=220)
        arq_url = st.text_input("ArquivoURL (link opcional — Drive/Cloudinary)", "")

        ok_save = st.form_submit_button("Salvar relatório")

    if ok_save:
        if not titulo.strip():
            st.error("Informe o título.")
        else:
            rid = new_id("R")
            row = {
                "RelatorioID": rid,
                "PacienteID": pid,
                "Data": data_rel.strftime(DATA_FMT),
                "Tipo": tipo,
                "Titulo": titulo.strip(),
                "Autor": (autor or "").strip() or "Equipe",
                "Texto": (texto or "").strip(),
                "ArquivoURL": (arq_url or "").strip(),
            }
            append_rows(ws_rel, [row], default_headers=REL_COLS)
            st.success(f"Relatório salvo ({rid}).")
            st.cache_data.clear()
            st.rerun()

# ---------- Sessões ----------
with tab_ses:
    st.subheader("Sessões do paciente")
    if df_ses_p.empty:
        st.info("Sem sessões.")
    else:
        show_cols = ["Data","HoraInicio","HoraFim","Profissional","Status","Tipo","ObjetivosTrabalhados","Observacoes"]
        show_cols = [c for c in show_cols if c in df_ses_p.columns]
        df_ses_p["__ord_h"] = pd.to_datetime(df_ses_p.get("HoraInicio",""), format="%H:%M", errors="coerce")
        df_ses_p2 = df_ses_p.sort_values(["__dt","__ord_h"], ascending=[True, True])
        st.dataframe(df_ses_p2[show_cols], use_container_width=True, hide_index=True)

# ---------- Financeiro ----------
with tab_fin:
    st.subheader("Financeiro (recebido)")
    if df_pag_p.empty:
        st.info("Sem pagamentos.")
    else:
        show_cols = ["Data","Forma","Bruto","Liquido","TaxaValor","Referencia","Obs"]
        show_cols = [c for c in show_cols if c in df_pag_p.columns]
        st.dataframe(df_pag_p.sort_values("__dt")[show_cols], use_container_width=True, hide_index=True)
        st.metric("Total líquido deste paciente", brl(float(df_pag_p['__liq'].sum())))

# ---------- Documentos ----------
with tab_docs:
    st.subheader("Documentos & anexos")
    st.info("Para anexar um link/arquivo, use o campo **ArquivoURL** ao criar um Relatório acima.\n\n"
            "Se preferir uma aba dedicada 'Documentos' na planilha (ex.: colunas: PacienteID, Titulo, Data, URL, Obs), dá pra adicionar depois.")
    rel_com_link = df_rel_p[df_rel_p.get("ArquivoURL","").astype(str).str.strip() != ""].copy()
    if rel_com_link.empty:
        st.caption("Nenhum link anexado ainda (ArquivoURL está vazio nos relatórios).")
    else:
        rel_com_link["__dt"] = rel_com_link.get("Data","").apply(to_date)
        rel_com_link = rel_com_link.sort_values("__dt")
        for _, r in rel_com_link.iterrows():
            dtxt = (to_date(r.get("Data")).strftime(DATA_FMT) if to_date(r.get("Data")) else "-")
            url = str(r.get("ArquivoURL")).strip()
            titulo = (str(r.get("Titulo","")).strip() or "Documento")
            autor = (str(r.get("Autor","")).strip() or nome_sel)
            st.markdown(f"- **{dtxt}** — {autor} · [{titulo}]({url})")
